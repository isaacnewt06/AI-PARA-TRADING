"""Extraction pipeline for PDFs, DOCX and XLSX files."""

from __future__ import annotations

from datetime import datetime, timezone
from pathlib import Path

import fitz
import openpyxl
from docx import Document as DocxDocument
from sqlalchemy.orm import Session

from src.core.config import Settings
from src.core.logging import get_logger
from src.db.models.document import Document
from src.db.models.file_asset import FileAsset
from src.processing.text_cleaner import TextCleaner
from src.telegram.parsers import TelegramMessageParser

logger = get_logger(__name__)


class DocumentProcessor:
    """Extract text from supported office documents."""

    def __init__(self, session: Session, settings: Settings) -> None:
        self.session = session
        self.settings = settings

    def process(self, file_asset: FileAsset) -> Document:
        path = Path(file_asset.stored_path)
        if not path.exists():
            raise FileNotFoundError(f"Document file not found: {file_asset.file_name}")
        if path.stat().st_size <= 0:
            raise ValueError(f"Empty document file: {file_asset.file_name}")
        if not self.is_supported(path, file_asset.file_name):
            raise ValueError(f"Unsupported document type for processing: {file_asset.file_name}")
        suffix = path.suffix.lower()
        if suffix == ".pdf":
            extracted = self._extract_pdf(path)
            doc_type = "pdf"
        elif suffix == ".docx":
            extracted = self._extract_docx(path)
            doc_type = "docx"
        elif suffix in {".xlsx", ".xlsm"}:
            extracted = self._extract_xlsx(path)
            doc_type = "xlsx"
        else:
            extracted = path.read_text(encoding="utf-8", errors="ignore")
            doc_type = "text"

        summary = self._local_summary(extracted)
        out_dir = self.settings.paths.processed_dir / "documents"
        out_dir.mkdir(parents=True, exist_ok=True)
        out_path = out_dir / f"file_{file_asset.id}.txt"
        out_path.write_text(extracted, encoding="utf-8")

        document = file_asset.document or Document(file_id=file_asset.id, doc_type=doc_type)
        document.extracted_text = extracted
        document.extracted_text_path = str(out_path.resolve())
        document.section_index_json = self._detect_sections(extracted)
        document.summary = summary
        document.processed_at = datetime.now(timezone.utc)
        self.session.add(document)
        file_asset.status = "processed"
        self.session.flush()
        logger.info("Processed document %s", file_asset.file_name)
        return document

    @staticmethod
    def _extract_pdf(path: Path) -> str:
        pages: list[str] = []
        with fitz.open(path) as document:
            for page_number, page in enumerate(document, start=1):
                text = page.get_text("text")
                pages.append(f"# Page {page_number}\n{text}")
        return TextCleaner.clean("\n\n".join(pages))

    @staticmethod
    def _extract_docx(path: Path) -> str:
        document = DocxDocument(path)
        paragraphs = [paragraph.text for paragraph in document.paragraphs if paragraph.text.strip()]
        return TextCleaner.clean("\n".join(paragraphs))

    @staticmethod
    def _extract_xlsx(path: Path) -> str:
        workbook = openpyxl.load_workbook(path, read_only=True, data_only=True)
        try:
            sheet_texts: list[str] = []
            for sheet in workbook.worksheets:
                rows: list[str] = []
                for row in sheet.iter_rows(values_only=True):
                    values = [str(value) for value in row if value is not None]
                    if values:
                        rows.append(" | ".join(values))
                if rows:
                    sheet_texts.append(f"# Sheet: {sheet.title}\n" + "\n".join(rows))
            return TextCleaner.clean("\n\n".join(sheet_texts))
        finally:
            workbook.close()

    @staticmethod
    def _local_summary(text: str) -> str:
        sentences = [segment.strip() for segment in text.split(".") if segment.strip()]
        return ". ".join(sentences[:3])[:600]

    @staticmethod
    def _detect_sections(text: str) -> str | None:
        lines = [line.strip() for line in text.splitlines() if line.strip()]
        sections = []
        for index, line in enumerate(lines):
            if line.startswith("#") or (len(line) < 80 and line == line.upper()):
                sections.append({"position": index, "title": line})
        if not sections:
            return None
        import json

        return json.dumps(sections, ensure_ascii=False)

    @staticmethod
    def is_supported(path: Path, file_name: str) -> bool:
        return TelegramMessageParser.is_supported_document_filename(file_name or path.name)
